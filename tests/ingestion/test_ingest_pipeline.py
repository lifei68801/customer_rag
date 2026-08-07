import aiosqlite

from app.graphrag.ontology import Term
from app.graphrag.review_queue import ensure_review_schema, list_pending_reviews
from app.ingestion.pipeline import (
    ingest_directory,
    ingest_docx_file,
    ingest_image_file,
    ingest_markdown_file,
    ingest_pdf_file,
    ingest_ticket_csv_file,
)
from app.providers.base import ProviderCapability, ProviderRequest, ProviderResult
from app.providers.embedding import EmbeddingRegistry, EmbeddingRequest, EmbeddingResult
from app.providers.registry import ProviderRegistry
from app.retrieval.vector_store import InMemoryVectorStore


class FakeEmbeddingProvider:
    async def embed(self, request: EmbeddingRequest) -> EmbeddingResult:
        return EmbeddingResult(vectors=[[0.1, 0.2] for _ in request.texts])


async def test_ingest_markdown_file_chunks_embeds_and_upserts(tmp_path):
    md_file = tmp_path / "network.md"
    md_file.write_text(
        "## 网络故障\n网络断开时请先重启路由器。\n"
        "## 登录问题\n登录失败请检查账号密码。\n",
        encoding="utf-8",
    )

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    count = await ingest_markdown_file(
        md_file,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
    )

    assert count == 2
    results = await vector_store.search(
        query_vector=[0.1, 0.2], top_k=2, tenant_id="t1"
    )
    assert len(results) == 2
    texts = {record.text for record in results}
    assert "网络断开时请先重启路由器。" in texts
    assert "登录失败请检查账号密码。" in texts


async def test_ingest_markdown_file_tags_records_with_the_ingesting_tenant(tmp_path):
    md_file = tmp_path / "network.md"
    md_file.write_text("## 网络故障\n网络断开时请先重启路由器。\n", encoding="utf-8")

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    await ingest_markdown_file(
        md_file,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
    )

    same_tenant = await vector_store.search(
        query_vector=[0.1, 0.2], top_k=10, tenant_id="t1"
    )
    other_tenant = await vector_store.search(
        query_vector=[0.1, 0.2], top_k=10, tenant_id="t2"
    )
    assert len(same_tenant) == 1
    assert len(other_tenant) == 0


async def test_ingest_markdown_file_skips_graph_extraction_when_not_configured(
    tmp_path,
):
    """向后兼容：不传图谱相关参数时行为与阶段2完全一致，不做任何图谱写入。"""
    md_file = tmp_path / "network.md"
    md_file.write_text("## 网络故障\n网络断开时请先重启路由器。\n", encoding="utf-8")

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    count = await ingest_markdown_file(
        md_file,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
    )

    assert count == 1


class FixedLLMProvider:
    def __init__(self, text: str) -> None:
        self._text = text

    async def complete(self, request: ProviderRequest) -> ProviderResult:
        return ProviderResult(text=self._text)


class FakeGraphClient:
    def __init__(self) -> None:
        self.written: list[dict] = []
        self.deleted_sources: list[tuple[str, str]] = []

    async def merge_relation(
        self,
        *,
        subject_standard_name,
        object_standard_name,
        relation_type,
        source,
        tenant_id,
    ) -> None:
        self.written.append(
            {
                "subject": subject_standard_name,
                "object": object_standard_name,
                "relation_type": relation_type,
                "source": source,
                "tenant_id": tenant_id,
            }
        )

    async def delete_relations_by_source(self, source: str, *, tenant_id: str) -> None:
        self.deleted_sources.append((source, tenant_id))


async def test_ingest_markdown_file_writes_graph_relations_when_configured(tmp_path):
    md_file = tmp_path / "network.md"
    md_file.write_text(
        "## 网络故障\n网关超时示例通常与示例认证模块相关\n", encoding="utf-8"
    )

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    llm_registry = ProviderRegistry()
    llm_registry.register(
        ProviderCapability.LLM,
        "llm",
        FixedLLMProvider(
            '{"relations": [{"subject": "网关超时示例", '
            '"object": "示例认证模块", "relation_type": "RELATED_TO"}]}'
        ),
    )
    terms = [
        Term(
            standard_name="示例错误码E502",
            aliases=["网关超时示例"],
            term_type="error_code",
            product_line="示例产品线",
        ),
        Term(
            standard_name="示例登录模块",
            aliases=["示例认证模块"],
            term_type="module",
            product_line="示例产品线",
        ),
    ]
    graph_client = FakeGraphClient()

    await ingest_markdown_file(
        md_file,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
        graph_llm_registry=llm_registry,
        graph_llm_provider_name="llm",
        graph_terms=terms,
        graph_client=graph_client,
    )

    assert graph_client.written == [
        {
            "subject": "示例错误码E502",
            "object": "示例登录模块",
            "relation_type": "RELATED_TO",
            "source": str(md_file),
            "tenant_id": "t1",
        }
    ]


async def test_ingest_markdown_file_sends_unresolved_candidates_to_review_queue(
    tmp_path,
):
    md_file = tmp_path / "network.md"
    md_file.write_text(
        "## 网络故障\n网关超时示例通常与不存在的实体相关\n", encoding="utf-8"
    )

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    llm_registry = ProviderRegistry()
    llm_registry.register(
        ProviderCapability.LLM,
        "llm",
        FixedLLMProvider(
            '{"relations": [{"subject": "网关超时示例", '
            '"object": "不存在的实体", "relation_type": "RELATED_TO"}]}'
        ),
    )
    terms = [
        Term(
            standard_name="示例错误码E502",
            aliases=["网关超时示例"],
            term_type="error_code",
            product_line="示例产品线",
        ),
    ]
    graph_client = FakeGraphClient()
    review_conn = await aiosqlite.connect(":memory:")
    await ensure_review_schema(review_conn)

    await ingest_markdown_file(
        md_file,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
        graph_llm_registry=llm_registry,
        graph_llm_provider_name="llm",
        graph_terms=terms,
        graph_client=graph_client,
        graph_review_conn=review_conn,
    )

    assert graph_client.written == []
    pending = await list_pending_reviews(review_conn)
    assert len(pending) == 1
    assert pending[0]["object_candidate"] == "不存在的实体"


async def test_ingest_docx_file_chunks_embeds_and_upserts(tmp_path):
    from docx import Document

    docx_path = tmp_path / "manual.docx"
    doc = Document()
    doc.add_heading("安装步骤", level=1)
    doc.add_paragraph("先下载安装包。")
    doc.save(str(docx_path))

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    count = await ingest_docx_file(
        docx_path,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
    )

    assert count == 1
    results = await vector_store.search(
        query_vector=[0.1, 0.2], top_k=1, tenant_id="t1"
    )
    assert "先下载安装包" in results[0].text


async def test_ingest_image_file_uses_injected_ocr_function(tmp_path):
    image_path = tmp_path / "scan.png"
    image_path.write_bytes(b"fake image bytes")

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    count = await ingest_image_file(
        image_path,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
        ocr=lambda path: "错误码E502表示网关超时",
    )

    assert count == 1
    results = await vector_store.search(
        query_vector=[0.1, 0.2], top_k=1, tenant_id="t1"
    )
    assert "错误码E502表示网关超时" in results[0].text


async def test_ingest_pdf_file_uses_injected_ocr_for_scanned_pages(tmp_path):
    from PIL import Image
    from reportlab.pdfgen import canvas

    scan_image_path = tmp_path / "scan.png"
    Image.new("RGB", (100, 100), color="white").save(scan_image_path)
    pdf_path = tmp_path / "scanned.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.drawImage(str(scan_image_path), 100, 700, width=100, height=100)
    c.showPage()
    c.save()

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    count = await ingest_pdf_file(
        pdf_path,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
        ocr=lambda path: "扫描件识别出的文字",
    )

    assert count == 1
    results = await vector_store.search(
        query_vector=[0.1, 0.2], top_k=1, tenant_id="t1"
    )
    assert "扫描件识别出的文字" in results[0].text


async def test_ingest_pdf_file_stores_full_table_text_for_table_row_chunks(tmp_path):
    # embedding 用的是每一行的细粒度文本（parent-child 的 child），但真正
    # 写进向量库、将来命中后返回给 LLM 的应该是整张表（parent），不能只有
    # 命中的那一行、丢了表格其余部分的上下文。
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

    if "STSong-Light" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    pdf_path = tmp_path / "table.pdf"
    doc = SimpleDocTemplate(str(pdf_path))
    table = Table(
        [
            ["股东名称", "持股比例"],
            ["武汉金融控股", "75.00%"],
            ["东亚银行", "15.38%"],
        ]
    )
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
            ]
        )
    )
    doc.build([table])

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    await ingest_pdf_file(
        pdf_path,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
    )

    all_records = await vector_store.list_all()
    table_records = [
        r for r in all_records if r.metadata.get("heading_path") == "第1页/表格"
    ]
    assert len(table_records) == 2
    assert all(
        "武汉金融控股" in r.text and "东亚银行" in r.text for r in table_records
    )


async def test_ingest_ticket_csv_file_chunks_embeds_and_upserts(tmp_path):
    import csv

    csv_path = tmp_path / "tickets.csv"
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=["ticket_id", "subject", "resolution"])
        writer.writeheader()
        writer.writerow(
            {"ticket_id": "T1", "subject": "登录失败", "resolution": "清除缓存后重新登录。"}
        )

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    count = await ingest_ticket_csv_file(
        csv_path,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
    )

    assert count == 1
    results = await vector_store.search(
        query_vector=[0.1, 0.2], top_k=1, tenant_id="t1"
    )
    assert "登录失败" in results[0].text
    assert "清除缓存后重新登录。" in results[0].text


async def test_ingest_directory_processes_markdown_and_pdf_but_skips_other_extensions(
    tmp_path,
):
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    (tmp_path / "a.md").write_text(
        "## 主题A\n内容A。\n", encoding="utf-8"
    )
    (tmp_path / "notes.txt").write_text("不应被处理", encoding="utf-8")

    pdf_path = tmp_path / "b.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.setFont("STSong-Light", 12)
    c.drawString(100, 750, "内容B。")
    c.showPage()
    c.save()

    embedding_registry = EmbeddingRegistry()
    embedding_registry.register("fake-embedding", FakeEmbeddingProvider())
    vector_store = InMemoryVectorStore()

    total = await ingest_directory(
        tmp_path,
        embedding_registry=embedding_registry,
        embedding_provider_name="fake-embedding",
        vector_store=vector_store,
        tenant_id="t1",
    )

    assert total == 2
    results = await vector_store.search(
        query_vector=[0.1, 0.2], top_k=10, tenant_id="t1"
    )
    texts = {record.text for record in results}
    assert texts == {"内容A。", "内容B。"}
