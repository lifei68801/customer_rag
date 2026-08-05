from docx import Document

from app.ingestion.docx_parser import parse_docx


def _write_docx(path, sections: list[tuple[str, str]]) -> None:
    doc = Document()
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    doc.save(str(path))


def test_parse_docx_returns_one_chunk_per_heading_section(tmp_path):
    docx_path = tmp_path / "manual.docx"
    _write_docx(
        docx_path,
        [("安装步骤", "先下载安装包。"), ("常见故障", "尝试重启设备。")],
    )

    chunks = parse_docx(docx_path)

    assert len(chunks) == 2
    assert chunks[0].heading_path == ["安装步骤"]
    assert "先下载安装包" in chunks[0].text
    assert chunks[0].source == str(docx_path)
    assert chunks[1].heading_path == ["常见故障"]
    assert "尝试重启设备" in chunks[1].text


def test_parse_docx_without_headings_returns_single_chunk(tmp_path):
    docx_path = tmp_path / "no_heading.docx"
    doc = Document()
    doc.add_paragraph("没有标题的正文内容。")
    doc.save(str(docx_path))

    chunks = parse_docx(docx_path)

    assert len(chunks) == 1
    assert chunks[0].heading_path == []
    assert "没有标题的正文内容" in chunks[0].text


def test_parse_docx_skips_empty_sections(tmp_path):
    docx_path = tmp_path / "empty_section.docx"
    doc = Document()
    doc.add_heading("空章节", level=1)
    doc.add_heading("有内容的章节", level=1)
    doc.add_paragraph("这里有内容。")
    doc.save(str(docx_path))

    chunks = parse_docx(docx_path)

    assert len(chunks) == 1
    assert chunks[0].heading_path == ["有内容的章节"]
